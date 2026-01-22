import os
import numpy as np
from docx import Document
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings

# -----------------------------
# SBERT 本地模型
# -----------------------------
MODEL_PATH = r"E:/GraduationProject/model/paraphrase-multilingual-MiniLM-L12-v2"
model = SentenceTransformer(MODEL_PATH, local_files_only=True)

# -----------------------------
# Chroma 0.3.4 初始化
# -----------------------------
client = chromadb.Client(Settings(
    chroma_db_impl="duckdb+parquet",
    persist_directory="vector_db"
))

# -----------------------------
# Collection 辅助函数
# -----------------------------
def get_or_create_collection(name):
    existing = [c.name for c in client.list_collections()]
    if name in existing:
        return client.get_collection(name)
    else:
        return client.create_collection(name)

# 创建/获取 Collection
collections = {
    "scenic": get_or_create_collection("scenic_knowledge"),
    "route": get_or_create_collection("route_knowledge"),
    "service": get_or_create_collection("service_knowledge")
}

# -----------------------------
# 文本处理函数
# -----------------------------
def load_docx_text(path):
    """加载 docx 文本，返回段落列表"""
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def chunk_text(texts, max_len=200):
    """按长度切分文本，返回文本块列表"""
    chunks = []
    for t in texts:
        if len(t) <= max_len:
            chunks.append(t)
        else:
            for i in range(0, len(t), max_len):
                chunks.append(t[i:i+max_len])
    return chunks

def chunk_scenic_docx_by_spot(path):
    """
    按景点完整信息分块：
    每个景点包含：
    建筑名、位置、开放时间、核心展品、建筑简介
    """
    doc = Document(path)
    chunks = []
    current_spot = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            if current_spot:
                chunks.append("\n".join(current_spot))
                current_spot = []
            continue
        current_spot.append(text)

    if current_spot:
        chunks.append("\n".join(current_spot))

    return chunks

# -----------------------------
# 构建向量库
# -----------------------------
import os

def vector_db_exists():
    """
    判断给定模式的向量库是否存在
    """
    persist_dir = os.path.join("vector_db", "index")
    # 检查 index 文件夹内是否有文件
    return os.path.exists(persist_dir) and len(os.listdir(persist_dir)) == 12

def build_vector_store(docx_path, mode):
    col = collections[mode]

    # 判断向量库是否已存在
    if vector_db_exists():
        print(f"向量库已存在，跳过构建过程")
        return  # 如果存在，直接跳过构建

    # 否则，构建向量库
    print(f"向量库不存在，开始构建")

    if mode == "scenic":
        chunks = chunk_scenic_docx_by_spot(docx_path)
    if mode == "service":
        chunks = chunk_service_docx_by_section(docx_path)
    if mode == "route":
        # ★ route 不切块，整个文档一条
        texts = load_docx_text(docx_path)
        chunks = ["\n".join(texts)]
    else:
        texts = load_docx_text(docx_path)

    embeddings = model.encode(chunks, convert_to_numpy=True)

    col.add(
        documents=chunks,
        embeddings=embeddings.tolist(),
        ids=[f"{mode}_chunk_{i}" for i in range(len(chunks))]
    )

    client.persist()
    print(f"[{mode}] 向量库构建完成，共 {len(chunks)} 个向量")

# -----------------------------
# 查询向量库
# -----------------------------
def query_vector_store(query_text, mode, top_k=7):
    query_emb = model.encode([query_text], convert_to_numpy=True)
    col = collections[mode]
    if mode == "scenic":
        top_k= 30

    results = col.query(
        query_embeddings=query_emb.tolist(),
        n_results=top_k
    )

    docs = results.get("documents", [[]])[0]

    # -------- 关键修改点 2 --------
    # service：保持原样，完整返回（不动）
    if mode == "service":
        if not docs:
            # 兜底：返回所有 service chunks（数量一般不大）
            return col.get()["documents"]
        return docs

    # -------- 关键修改点 3 --------
    # scenic：只返回一个景点，且限制长度（建议 400 以内）
    if mode == "scenic":
        return docs  # 返回 top_k 个景点 chunk

def chunk_service_docx_by_section(path):
    """
    按“数字编号标题”切分基础服务文档
    例如：1. 入馆路线与通行规则
    """
    doc = Document(path)

    chunks = []
    current_title = None
    current_lines = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        # 匹配一级标题，如 "1. 入馆路线与通行规则"
        if text[0].isdigit() and "." in text[:3]:
            if current_title and current_lines:
                chunks.append(
                    current_title + "\n" + "\n".join(current_lines)
                )
            current_title = text
            current_lines = []
        else:
            current_lines.append(text)

    # 最后一块
    if current_title and current_lines:
        chunks.append(
            current_title + "\n" + "\n".join(current_lines)
        )

    return chunks

# -----------------------------
# 对外统一搜索接口（给 main.py 用）
# -----------------------------
def search(query: str, mode: str, top_k: int = 5):
    """
    对外统一向量搜索接口
    返回：List[str]
    """
    return query_vector_store(
        query_text=query,
        mode=mode,
        top_k=top_k
    )

# -----------------------------
# 初始化时生成三个模式向量库
# -----------------------------
DOCX_FILES = {
    "scenic": "knowledge/景点介绍.docx",
    "route": "knowledge/预设路线.docx",
    "service": "knowledge/基础实用服务信息.docx"
}

for mode, path in DOCX_FILES.items():
    if os.path.exists(path):
        build_vector_store(path, mode)
    else:
        print(f"[警告] {path} 不存在，跳过 {mode} 向量库构建")
